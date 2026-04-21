"""
DocFlow Pro — Servidor principal
Convierte PDF a Word con OCR inteligente
"""
import os
import uuid
import shutil
import subprocess
import threading
import tempfile
import time
from pathlib import Path
from flask import Flask, request, jsonify, send_file, send_from_directory
from werkzeug.utils import secure_filename

app = Flask(__name__, static_folder='/app/static')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB

WORK_DIR = Path(tempfile.gettempdir()) / 'docflow'
WORK_DIR.mkdir(exist_ok=True)

jobs = {}
jobs_lock = threading.Lock()
JOB_TTL = 3600

# ─── Límite de conversiones por IP ───────────────────────────────────────────
DAILY_LIMIT = 3
ip_usage = {}
ip_lock = threading.Lock()

def get_client_ip():
    if request.headers.get('X-Forwarded-For'):
        return request.headers.get('X-Forwarded-For').split(',')[0].strip()
    return request.remote_addr

def get_today():
    import datetime
    return datetime.date.today().isoformat()

def check_limit(ip):
    today = get_today()
    with ip_lock:
        data = ip_usage.get(ip)
        if not data or data['date'] != today:
            ip_usage[ip] = {'count': 0, 'date': today}
            data = ip_usage[ip]
        remaining = DAILY_LIMIT - data['count']
        return remaining > 0, remaining

def consume_limit(ip):
    today = get_today()
    with ip_lock:
        if ip not in ip_usage or ip_usage[ip]['date'] != today:
            ip_usage[ip] = {'count': 0, 'date': today}
        ip_usage[ip]['count'] += 1

def cleanup_ip_usage():
    while True:
        time.sleep(3600)
        today = get_today()
        with ip_lock:
            stale = [ip for ip, d in ip_usage.items() if d['date'] != today]
            for ip in stale:
                del ip_usage[ip]

threading.Thread(target=cleanup_ip_usage, daemon=True).start()

def cleanup_old_jobs():
    while True:
        time.sleep(600)
        now = time.time()
        to_delete = []
        with jobs_lock:
            for jid, job in jobs.items():
                if now - job.get('created_at', now) > JOB_TTL:
                    to_delete.append(jid)
            for jid in to_delete:
                job = jobs.pop(jid)
                out = job.get('output_file')
                if out:
                    Path(out).unlink(missing_ok=True)

threading.Thread(target=cleanup_old_jobs, daemon=True).start()


# ─── Detección de PDF escaneado ───────────────────────────────────────────────

def is_scanned_pdf(pdf_path: Path) -> bool:
    """
    Detecta si el PDF es escaneado (imagen) o tiene texto real.
    Usa pdfminer para extraer texto — si hay muy poco, es escaneado.
    """
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(str(pdf_path))
        # Si tiene menos de 50 caracteres reales, es escaneado
        clean = text.strip().replace('\n', '').replace(' ', '')
        return len(clean) < 50
    except Exception:
        return True  # Si falla, asumir escaneado y aplicar OCR


# ─── Conversión con OCR ───────────────────────────────────────────────────────

def convert_with_ocr(pdf_path: Path, job_id: str) -> Path:
    """
    Convierte PDF escaneado a DOCX usando OCR:
    1. PDF → imágenes (pdf2image)
    2. Imágenes → texto (Tesseract OCR)
    3. Texto → DOCX (python-docx)
    """
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image

    update_job(job_id, step='Convirtiendo páginas a imágenes...', progress=40)

    # Convertir PDF a imágenes
    images = convert_from_path(
        str(pdf_path),
        dpi=300,
        fmt='PNG',
        output_folder=str(WORK_DIR),
        output_file=f'{job_id}_page'
    )

    update_job(job_id, step=f'Aplicando OCR a {len(images)} página(s)...', progress=55)

    # Crear DOCX con python-docx
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Si python-docx no está instalado, instalar en tiempo de ejecución
        subprocess.run(['pip', 'install', 'python-docx', '--break-system-packages', '-q'],
                      capture_output=True)
        from docx import Document
        from docx.shared import Pt, Inches

    doc = Document()

    # Configurar márgenes del documento
    from docx.shared import Inches
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    total_pages = len(images)
    for i, image in enumerate(images):
        update_job(job_id,
                   step=f'Procesando página {i+1} de {total_pages}...',
                   progress=55 + int((i / total_pages) * 30))

        # OCR con Tesseract — español primero, luego inglés
        try:
            text = pytesseract.image_to_string(
                image,
                lang='spa+eng',
                config='--psm 6 --oem 3'
            )
        except Exception:
            text = pytesseract.image_to_string(image, lang='eng')

        # Agregar texto al documento
        if text.strip():
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    para = doc.add_paragraph(line)
                    para.style.font.size = Pt(11)
                else:
                    doc.add_paragraph('')

        # Salto de página entre páginas del PDF (excepto la última)
        if i < total_pages - 1:
            doc.add_page_break()

    # Guardar DOCX
    output_path = WORK_DIR / f'{job_id}_ocr.docx'
    doc.save(str(output_path))

    # Limpiar imágenes temporales
    for image in images:
        try:
            image_path = Path(str(image.filename)) if hasattr(image, 'filename') else None
            if image_path and image_path.exists():
                image_path.unlink()
        except Exception:
            pass

    return output_path


# ─── Utilidades LibreOffice ───────────────────────────────────────────────────

def get_soffice_env():
    env = os.environ.copy()
    env['SAL_USE_VCLPLUGIN'] = 'svp'
    env['HOME'] = '/tmp'
    return env

def run_libreoffice(args, timeout=120):
    env = get_soffice_env()
    cmd = ['soffice'] + args
    return subprocess.run(cmd, env=env, capture_output=True, text=True, timeout=timeout)


# ─── Remoción de seguridad ────────────────────────────────────────────────────

def remove_security(pdf_path: Path, password: str = '') -> tuple:
    out = pdf_path.parent / (pdf_path.stem + '_open.pdf')
    out.unlink(missing_ok=True)

    candidates = []
    if password:
        candidates.append(password)
    candidates += ['', 'owner', 'user', '1234', 'password', 'pdf', 'admin', '0000']

    for pwd in candidates:
        cmd = ['qpdf']
        if pwd:
            cmd += [f'--password={pwd}']
        cmd += ['--decrypt', str(pdf_path), str(out)]
        try:
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if r.returncode == 0 and out.exists() and out.stat().st_size > 0:
                msg = 'Contrasena removida' if pwd else 'Sin contrasena detectada'
                return out, msg
        except Exception:
            pass
        out.unlink(missing_ok=True)

    try:
        cmd = ['qpdf', '--qdf', '--object-streams=disable', str(pdf_path), str(out)]
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if r.returncode == 0 and out.exists() and out.stat().st_size > 0:
            return out, 'Restricciones removidas'
    except Exception:
        pass

    return pdf_path, 'PDF procesado'


# ─── Conversión principal ─────────────────────────────────────────────────────

def update_job(job_id: str, **kwargs):
    with jobs_lock:
        if job_id in jobs:
            jobs[job_id].update(kwargs)

def do_convert(pdf_path: Path, job_id: str):
    clean_pdf = None
    try:
        update_job(job_id, step='Analizando PDF...', progress=10)

        with jobs_lock:
            password = jobs[job_id].get('password', '')

        # Paso 1: remover seguridad
        clean_pdf, security_msg = remove_security(pdf_path, password)
        update_job(job_id, step=f'{security_msg} — Analizando contenido...', progress=25)

        # Paso 2: detectar si es escaneado
        scanned = is_scanned_pdf(clean_pdf)
        final = WORK_DIR / f'{job_id}.docx'

        if scanned:
            # PDF escaneado → OCR
            update_job(job_id,
                      step='PDF escaneado detectado — Aplicando OCR inteligente...',
                      progress=35)
            ocr_output = convert_with_ocr(clean_pdf, job_id)
            shutil.move(str(ocr_output), str(final))
            method = 'OCR'
        else:
            # PDF con texto real → LibreOffice
            update_job(job_id, step='Texto digital detectado — Convirtiendo...', progress=35)

            result = run_libreoffice([
                '--headless',
                '--infilter=writer_pdf_import',
                '--convert-to', 'docx:MS Word 2007 XML',
                '--outdir', str(WORK_DIR),
                str(clean_pdf)
            ], timeout=150)

            update_job(job_id, progress=75)
            docx_file = WORK_DIR / (clean_pdf.stem + '.docx')

            if not docx_file.exists() or docx_file.stat().st_size == 0:
                # Reintento sin filtro
                update_job(job_id, step='Reintentando...', progress=80)
                run_libreoffice([
                    '--headless', '--convert-to', 'docx',
                    '--outdir', str(WORK_DIR), str(clean_pdf)
                ], timeout=150)
                docx_file = WORK_DIR / (clean_pdf.stem + '.docx')

            if not docx_file.exists() or docx_file.stat().st_size == 0:
                # Fallback: usar OCR igual
                update_job(job_id, step='Aplicando OCR como alternativa...', progress=80)
                ocr_output = convert_with_ocr(clean_pdf, job_id)
                shutil.move(str(ocr_output), str(final))
                method = 'OCR (fallback)'
            else:
                shutil.move(str(docx_file), str(final))
                method = 'conversión directa'

        update_job(job_id,
                  status='done',
                  step='Conversion completada!',
                  progress=100,
                  output_file=str(final),
                  security_msg=security_msg,
                  method=method)

    except subprocess.TimeoutExpired:
        update_job(job_id, status='error', error='Tiempo agotado. PDF muy grande o complejo.')
    except Exception as e:
        update_job(job_id, status='error', error=str(e))
    finally:
        try:
            pdf_path.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            if clean_pdf and clean_pdf != pdf_path:
                clean_pdf.unlink(missing_ok=True)
        except Exception:
            pass


# ─── Rutas Flask ──────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return send_from_directory('/app/static', 'index.html')

@app.route('/limit')
def limit_status():
    ip = get_client_ip()
    can_convert, remaining = check_limit(ip)
    return jsonify({'can_convert': can_convert, 'remaining': remaining, 'daily_limit': DAILY_LIMIT})

@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return jsonify({'error': 'No se recibio archivo'}), 400
    f = request.files['file']
    if not f.filename or not f.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Solo se aceptan archivos PDF'}), 400

    ip = get_client_ip()
    can_convert, remaining = check_limit(ip)
    if not can_convert:
        return jsonify({
            'error': 'limite_alcanzado',
            'message': 'Has alcanzado el limite de 3 conversiones gratuitas por hoy. Vuelve manana.'
        }), 429

    job_id = uuid.uuid4().hex[:10]
    safe_name = secure_filename(f.filename)
    upload_path = WORK_DIR / f'{job_id}_{safe_name}'
    f.save(str(upload_path))
    password = request.form.get('password', '')
    consume_limit(ip)

    with jobs_lock:
        jobs[job_id] = {
            'status': 'processing', 'step': 'Iniciando...', 'progress': 5,
            'original_name': safe_name, 'password': password,
            'output_file': None, 'created_at': time.time(),
        }

    t = threading.Thread(target=do_convert, args=(upload_path, job_id), daemon=True)
    t.start()
    return jsonify({'job_id': job_id, 'filename': safe_name, 'remaining': remaining - 1})

@app.route('/status/<job_id>')
def status(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job:
        return jsonify({'error': 'Trabajo no encontrado'}), 404
    safe = {k: v for k, v in job.items() if k not in ('output_file', 'password', 'created_at')}
    return jsonify(safe)

@app.route('/download/<job_id>')
def download(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job or job.get('status') != 'done':
        return jsonify({'error': 'No listo aun'}), 404
    out_file = Path(job['output_file'])
    if not out_file.exists():
        return jsonify({'error': 'Archivo no encontrado'}), 404
    original = Path(job.get('original_name', 'documento.pdf'))
    dl_name = original.stem + '_convertido.docx'
    return send_file(str(out_file), as_attachment=True, download_name=dl_name)

@app.route('/health')
def health():
    return jsonify({'status': 'ok', 'service': 'DocFlow Pro'})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port, debug=False, threaded=True)
